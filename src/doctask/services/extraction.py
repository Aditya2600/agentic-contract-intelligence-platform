"""Turn uploaded bytes into blocks a citation can point at.

Native first: PyMuPDF for PDF, python-docx for DOCX, a UTF-8 decode for TXT. A model is
only asked to read a page when the native extractor did not produce usable text for it,
so the normal path costs nothing and stays byte-faithful.

The rule that shapes everything here: a page this module cannot read is an error, never
an empty string. Silently dropping an unreadable page is how a scanned liability cap
becomes a `pass` or an `insufficient_evidence` that nobody questions -- the register
would look clean because the evidence was never in front of the model, not because the
contract was fine. So every failure raises `ExtractionError` and the run never starts.
"""

from __future__ import annotations

import base64
import re
from dataclasses import dataclass, field
from time import perf_counter
from typing import Protocol

# Minimum alphanumeric characters before a page counts as read at all.
_MIN_ALNUM = 30
# Below this share of letters among non-space characters, the text is mojibake or a
# table of glyph soup rather than prose.
_MIN_ALPHA_RATIO = 0.35
# A page that is mostly picture and hardly any text is a scan, whatever it decodes to.
_VISUAL_COVERAGE = 0.5
_VISUAL_TEXT_LIMIT = 400
_OCR_DPI = 200
# Horizontal gap, in points, that no text block crosses before it counts as a column
# boundary rather than ordinary word spacing.
_COLUMN_GAP = 20.0
_COLUMN_SLACK = 2.0
# Below this many blocks there is no layout to infer; treat the page as one column.
_MIN_COLUMN_BLOCKS = 4

NATIVE_PDF = "native_pdf"
GEMMA_VLM = "gemma_vlm"
DOCX = "docx"
TXT = "txt"


class ExtractionError(Exception):
    """The bytes could not be turned into text that is safe to reason about."""


class PageOCR(Protocol):
    async def read_page(self, image_png: bytes, *, page: int) -> str:
        """Transcribe one rendered page. Empty string means nothing legible."""
        ...


@dataclass(frozen=True, slots=True)
class ExtractedBlock:
    text: str
    page: int | None
    extraction_method: str


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    """`text` is exactly the blocks joined by blank lines, which is what makes a block's
    offsets into it recomputable and its quotes verbatim.

    `warnings` is everything about the reading that a reviewer would want to know before
    believing a verdict drawn from it: a page a model transcribed, a rotated page, a
    table flattened into rows, a file that was encrypted. None of these is bad enough to
    refuse the document, and every one of them is a reason a rule might have been judged
    on evidence that is not quite what the file says -- so they travel with the run and
    stop it claiming a clean result.
    """

    text: str
    blocks: list[ExtractedBlock]
    warnings: list[str] = field(default_factory=list)
    # One entry per page a vision model transcribed: what it cost, in the pipeline's own
    # units. Extraction runs before a run_id exists to attach a `RunEvent` to, so this is
    # how that cost survives to reach one -- `ingest_file` replays these into the run's
    # event log once the run it belongs to has been created.
    ocr_calls: list[dict] = field(default_factory=list)

    @property
    def methods(self) -> set[str]:
        return {block.extraction_method for block in self.blocks}


def _paragraphs(text: str) -> list[str]:
    """Split on blank lines, and collapse blank lines *inside* a paragraph.

    A block that still contained a blank line would be re-split downstream, and its
    stored offsets would then point at the wrong span.
    """
    parts = re.split(r"\n\s*\n", text)
    return [re.sub(r"\n\s*\n+", "\n", part).strip() for part in parts if part.strip()]


def _alnum(text: str) -> int:
    return sum(1 for char in text if char.isalnum())


def _is_garbled(text: str) -> bool:
    dense = [char for char in text if not char.isspace()]
    if not dense:
        return True
    if "�" in text:
        return True
    letters = sum(1 for char in dense if char.isalpha())
    return letters / len(dense) < _MIN_ALPHA_RATIO


def _needs_ocr(text: str, *, image_coverage: float, has_table: bool = False) -> str | None:
    """Why this page cannot be trusted as extracted, or None if it can."""
    if _alnum(text) < _MIN_ALNUM:
        return "native extraction produced no readable text"
    # A table is mostly digits and separators by nature. Judging it on letter density
    # would send a legible fee schedule to OCR, or fail the upload over a page that
    # extracted perfectly well.
    if not has_table and _is_garbled(text):
        return "native extraction produced garbled text"
    if image_coverage >= _VISUAL_COVERAGE and _alnum(text) < _VISUAL_TEXT_LIMIT:
        return "page is mostly image with little extractable text"
    return None


def _image_coverage(page) -> float:
    page_area = abs(page.rect.get_area())
    if page_area <= 0:
        return 0.0
    covered = 0.0
    for block in page.get_text("dict").get("blocks", []):
        if block.get("type") == 1:  # image block
            x0, y0, x1, y1 = block["bbox"]
            covered += abs((x1 - x0) * (y1 - y0))
    return min(covered / page_area, 1.0)


def _columns(boxes: list[tuple[float, float, float, float]], page_width: float) -> list[float]:
    """Left edges of the page's text columns, or a single column.

    Sorting purely by vertical position interleaves a two-column layout: the right
    column's first paragraph sits level with the left column's first paragraph, so
    reading order alternates between them and every sentence lands next to one it does
    not belong with. Columns are found as horizontal bands that no block straddles.
    """
    if len(boxes) < _MIN_COLUMN_BLOCKS:
        return [0.0]
    spans = sorted((x0, x1) for x0, _, x1, _ in boxes)
    boundaries: list[float] = [spans[0][0]]
    reach = spans[0][1]
    for x0, x1 in spans[1:]:
        if x0 > reach + _COLUMN_GAP:
            boundaries.append(x0)
            reach = x1
        else:
            reach = max(reach, x1)
    # A "column" narrower than a fifth of the page is a margin note or a stray box,
    # not a column, and splitting on it would scatter the body text.
    return boundaries if len(boundaries) > 1 and page_width > 0 else [0.0]


def _page_text(page) -> str:
    """Text blocks in reading order: column by column, top to bottom within each."""
    raw = [block for block in page.get_text("blocks") if block[4].strip()]
    if not raw:
        return ""
    boxes = [(block[0], block[1], block[2], block[3]) for block in raw]
    boundaries = _columns(boxes, page.rect.width)

    def column_of(x0: float) -> int:
        return max(index for index, left in enumerate(boundaries) if x0 >= left - _COLUMN_SLACK)

    ordered = sorted(raw, key=lambda block: (column_of(block[0]), round(block[1], 1), block[0]))
    return "\n\n".join(block[4].strip() for block in ordered)


def _table_text(page) -> str:
    """Tables rendered row by row, so a cell keeps the row it belongs to.

    A table read as loose text blocks comes out as a column of bare numbers with the
    headers somewhere above them, which is both unquotable and easy to misread. It also
    looks like mojibake to the garble check, which would send a perfectly legible
    financial table to OCR or fail the upload outright.
    """
    try:
        tables = page.find_tables()
    except Exception:  # pragma: no cover - pymupdf table finder is best effort
        return ""
    rows: list[str] = []
    for table in tables:
        for row in table.extract():
            cells = [str(cell).strip() for cell in row if cell and str(cell).strip()]
            if cells:
                rows.append(" | ".join(cells))
    return "\n".join(rows)


def _open_pdf(data: bytes):
    """Open a PDF, or say precisely why it cannot be read."""
    import pymupdf

    try:
        document = pymupdf.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise ExtractionError(f"unreadable PDF: {exc}") from exc

    locked = bool(document.needs_pass)
    if locked:
        # Many "encrypted" contracts are only owner-locked: an empty user password
        # opens them. One that genuinely needs a password is refused rather than
        # ingested as a zero-page document.
        if not document.authenticate(""):
            document.close()
            raise ExtractionError("PDF is password protected and no password was supplied")
    # `is_encrypted` reads False once MuPDF has decrypted the file, which it does with an
    # empty password before anything here runs. The metadata still names the scheme, and
    # that is the only thing left to tell a reviewer the file arrived encrypted.
    encrypted = locked or bool((document.metadata or {}).get("encryption"))
    if document.page_count == 0:
        document.close()
        raise ExtractionError("PDF has no pages")
    return document, encrypted, bool(getattr(document, "is_repaired", False))


async def _extract_pdf(
    data: bytes, ocr: PageOCR | None, warnings: list[str], ocr_calls: list[dict]
) -> list[ExtractedBlock]:
    try:
        import pymupdf  # noqa: F401
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise ExtractionError("PDF support needs pymupdf") from exc

    document, was_encrypted, was_repaired = _open_pdf(data)
    blocks: list[ExtractedBlock] = []
    with document:
        if was_encrypted:
            warnings.append("the file was encrypted and opened with an empty password")
        if was_repaired:
            # A damaged file MuPDF rebuilt. It opens and it reads, and there is no way
            # to know from the inside what the rebuild dropped.
            warnings.append("the file was damaged and had to be repaired before reading")
        for number, page in enumerate(document, start=1):
            try:
                text = _page_text(page)
                tables = _table_text(page)
                coverage = _image_coverage(page)
                rotation = page.rotation
            except Exception as exc:
                # A page that raises is a page nobody read. Skipping it would drop its
                # clauses from the corpus while the run still reported a result.
                raise ExtractionError(f"page {number} could not be read: {exc}") from exc

            if tables:
                # Table rows are appended, not substituted: the surrounding prose is
                # still the evidence for anything stated outside the grid.
                text = f"{text}\n\n{tables}".strip()
                warnings.append(f"page {number} contains a table, read cell by cell")

            problem = _needs_ocr(text, image_coverage=coverage, has_table=bool(tables))
            method = NATIVE_PDF
            if problem is not None:
                text = await _ocr_page(
                    page, number=number, ocr=ocr, problem=problem, ocr_calls=ocr_calls
                )
                method = GEMMA_VLM
                warnings.append(f"page {number} was read by a vision model, not from the file")
            elif rotation:
                warnings.append(f"page {number} is rotated {rotation} degrees")

            blocks.extend(
                ExtractedBlock(text=paragraph, page=number, extraction_method=method)
                for paragraph in _paragraphs(text)
            )
    if not blocks:
        raise ExtractionError("PDF produced no text blocks")
    return blocks


async def _ocr_page(
    page, *, number: int, ocr: PageOCR | None, problem: str, ocr_calls: list[dict]
) -> str:
    """Read a page the native extractor could not. Failing here fails the document."""
    if ocr is None:
        raise ExtractionError(f"page {number}: {problem}, and no OCR model is configured")
    image = page.get_pixmap(dpi=_OCR_DPI).tobytes("png")
    started = perf_counter()
    try:
        text = await ocr.read_page(image, page=number)
    except Exception as exc:
        raise ExtractionError(f"page {number}: {problem}, and OCR failed: {exc}") from exc
    usage = getattr(ocr, "last_usage", {}) or {}
    ocr_calls.append(
        {
            "page": number,
            "model": getattr(ocr, "vision_model", None) or getattr(ocr, "model", "unknown"),
            "tokens_in": usage.get("tokens_in", 0),
            "tokens_out": usage.get("tokens_out", 0),
            "duration_ms": int((perf_counter() - started) * 1000),
        }
    )
    if _alnum(text) < _MIN_ALNUM or _is_garbled(text):
        raise ExtractionError(f"page {number}: {problem}, and OCR returned nothing legible")
    return text


def _extract_docx(data: bytes, warnings: list[str]) -> list[ExtractedBlock]:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise ExtractionError("DOCX support needs python-docx") from exc

    from io import BytesIO

    try:
        document = docx.Document(BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"unreadable DOCX: {exc}") from exc

    parts = [paragraph.text.strip() for paragraph in document.paragraphs]
    if document.tables:
        warnings.append(f"{len(document.tables)} table(s) read cell by cell")
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    blocks = [
        ExtractedBlock(text=part, page=None, extraction_method=DOCX)
        for part in parts
        if part.strip()
    ]
    if not blocks:
        raise ExtractionError("DOCX produced no text blocks")
    return blocks


def _extract_txt(data: bytes) -> list[ExtractedBlock]:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError as exc:
        # Not decoded with errors="replace": a mangled clause is worse than a refusal,
        # because the mangled version still reads like evidence.
        raise ExtractionError(f"file is not valid UTF-8: {exc}") from exc
    blocks = [
        ExtractedBlock(text=paragraph, page=None, extraction_method=TXT)
        for paragraph in _paragraphs(text)
    ]
    if not blocks:
        raise ExtractionError("file contains no text")
    return blocks


def _kind(filename: str, mime_type: str) -> str:
    """PDF, DOCX or TXT, from the declared type or, failing that, the extension."""
    mime = (mime_type or "").split(";")[0].strip().lower()
    by_mime = {
        "application/pdf": "pdf",
        "application/x-pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "txt",
        "text/markdown": "txt",
    }
    if mime in by_mime:
        return by_mime[mime]
    suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    by_suffix = {"pdf": "pdf", "docx": "docx", "txt": "txt", "md": "txt", "text": "txt"}
    if suffix in by_suffix:
        return by_suffix[suffix]
    raise ExtractionError(f"unsupported document type: {mime_type or filename}")


async def extract_document(
    *, filename: str, mime_type: str, data: bytes, ocr: PageOCR | None = None
) -> ExtractedDocument:
    """Extract `data` into blocks, or raise. Never returns a partial document."""
    if not data:
        raise ExtractionError("uploaded file is empty")
    kind = _kind(filename, mime_type)
    warnings: list[str] = []
    ocr_calls: list[dict] = []
    if kind == "pdf":
        blocks = await _extract_pdf(data, ocr, warnings, ocr_calls)
    elif kind == "docx":
        blocks = _extract_docx(data, warnings)
    else:
        blocks = _extract_txt(data)
    return ExtractedDocument(
        text="\n\n".join(block.text for block in blocks),
        blocks=blocks,
        warnings=warnings,
        ocr_calls=ocr_calls,
    )


def png_data_url(image_png: bytes) -> str:
    """How a rendered page is handed to a vision model."""
    return "data:image/png;base64," + base64.b64encode(image_png).decode("ascii")
